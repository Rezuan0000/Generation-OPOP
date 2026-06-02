"""
Генерация документа ОПОП на основе шаблона Word и данных из JSON.

Перед использованием установите зависимость:
    pip install python-docx

В шаблоне используйте метки вида {{ключ}}, где ключ — это имя
поля из JSON-файла.
Например: {{direction_code}}, {{direction_name}}, {{profile_full}} и т.д.

Для динамических таблиц используйте маркеры:
    {{START_TABLE:universal}} / {{END_TABLE:universal}} - для таблицы УК
    {{START_TABLE:professional}} / {{END_TABLE:professional}} - для таблицы ОПК

Ручной ввод (значения в JSON: manual_fields["ключ"]):
    {{MANUAL:ключ}}  — например {{MANUAL:normative_docs}}

Подписи из БД (base64 в employee_signatures["ключ"]):
    {{IMAGE:ключ}}  — например {{IMAGE:first_responsible}}, {{IMAGE:second_responsible}}
"""

import base64
import json
import re
import zipfile
from io import BytesIO
from pathlib import Path
from copy import deepcopy
from typing import Any

from docx.shared import Pt, Cm
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn

from docx import Document
from zipfile import BadZipFile

# Ручной ввод в шаблоне: {{MANUAL:имя_поля}} (латиница, цифры, подчёркивание).
MANUAL_PLACEHOLDER_RE = re.compile(r"\{\{MANUAL:([a-zA-Z0-9_]+)\}\}")

# Подписи: {{IMAGE:имя}} — ключ совпадает с employee_signatures в opop_data.json.
IMAGE_PLACEHOLDER_RE = re.compile(r"\{\{IMAGE:([a-zA-Z0-9_]+)\}\}")

# Ширина вставляемой подписи по умолчанию (см).
_DEFAULT_SIGNATURE_WIDTH_CM = 3.5

# Символы, недопустимые в XML Word (python-docx).
_INVALID_XML_CHAR_RE = re.compile(
    r"[\x00-\x08\x0b\x0c\x0e-\x1f\ud800-\udfff\ufeff\ufffe\uffff]"
)


def _sanitize_xml_text(value: Any) -> str:
    """Убирает управляющие и прочие символы, из‑за которых падает запись в .docx."""
    if value is None:
        return ""
    if not isinstance(value, str):
        value = str(value)
    return _INVALID_XML_CHAR_RE.sub("", value)

# Ячейка с компетенцией из БД (УК/ОПК/ПК), без меток MANUAL — не очищаем при размножении строк.
_COMPETENCY_FROM_DB_RE = re.compile(r"^\s*(УК|ОПК|ПК)-\d+\s*-", re.IGNORECASE)


def _set_paragraph_text_preserve_first_run_style(paragraph, new_text: str) -> None:
    """
    Записывает новый текст в абзац и сохраняет базовый стиль первого run.
    Это уменьшает риск "скачка" шрифта после подстановки.
    """
    new_text = _sanitize_xml_text(new_text)
    runs = list(paragraph.runs)
    if not runs:
        paragraph.add_run(new_text)
        return

    first = runs[0]
    first.text = new_text
    for r in runs[1:]:
        r._element.getparent().remove(r._element)


def _normalize_manual_text(value: str) -> str:
    """Нормализует текст из textarea для вставки в Word."""
    normalized = _sanitize_xml_text(value).replace("\t", " ")
    lines = normalized.replace("\r\n", "\n").replace("\r", "\n").splitlines()
    lines = [line.lstrip() for line in lines]
    return "\n".join(lines)


def _set_paragraph_multiline_text_preserve_style(paragraph, new_text: str) -> None:
    """
    Записывает многострочный текст в исходный абзац через line break внутри run.
    Так сохраняется исходный стиль и исчезают "плавающие" отступы между строками.
    """
    runs = list(paragraph.runs)
    if not runs:
        run = paragraph.add_run()
    else:
        run = runs[0]
        for r in runs[1:]:
            r._element.getparent().remove(r._element)

    parts = new_text.split("\n")
    run.text = parts[0] if parts else ""
    for part in parts[1:]:
        run.add_break(WD_BREAK.LINE)
        run.add_text(part)


def _manual_lines(value: str) -> list[str]:
    normalized = _normalize_manual_text(value or "")
    lines = normalized.split("\n")
    return lines if lines else [""]


def _replace_manual_placeholders_in_row(table, row_idx: int, manual: dict[str, str], keys: set[str]) -> None:
    row = table.rows[row_idx]
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            if not paragraph.text:
                continue
            text = paragraph.text
            replaced = text
            changed = False
            for key in keys:
                placeholder = f"{{{{MANUAL:{key}}}}}"
                if placeholder in replaced:
                    replaced = replaced.replace(placeholder, _normalize_manual_text(manual.get(key, "")))
                    changed = True
            if changed:
                paragraph.paragraph_format.first_line_indent = Pt(0)
                _set_paragraph_multiline_text_preserve_style(paragraph, replaced)


def _cell_flat_text(cell) -> str:
    return " ".join(p.text for p in cell.paragraphs)


def _is_competency_cell_from_db(cell) -> bool:
    t = _cell_flat_text(cell).strip()
    if not t or MANUAL_PLACEHOLDER_RE.search(t):
        return False
    return bool(_COMPETENCY_FROM_DB_RE.match(t))


def _clear_cell_text_keep_paragraphs(cell) -> None:
    """Очищает текст в ячейке, не удаляя саму структуру ячейки/абзацев."""
    for paragraph in cell.paragraphs:
        if paragraph.runs:
            paragraph.runs[0].text = ""
            for run in paragraph.runs[1:]:
                run._element.getparent().remove(run._element)
        else:
            paragraph.add_run("")


def _expand_manual_rows_in_tables(doc: Document, manual: dict[str, str]) -> None:
    """
    Расширяет строки таблиц с {{MANUAL:key}}:
    - если в значениях есть несколько строк, создаёт несколько строк таблицы;
    - строка N берёт N-ю строку из каждого manual-поля (если нет — пусто).
    """
    for table in doc.tables:
        i = 0
        while i < len(table.rows):
            row = table.rows[i]
            row_text = " ".join(cell.text for cell in row.cells)
            matches = list(MANUAL_PLACEHOLDER_RE.finditer(row_text))
            if not matches:
                i += 1
                continue

            row_keys = {m.group(1) for m in matches}
            lines_by_key: dict[str, list[str]] = {k: _manual_lines(manual.get(k, "")) for k in row_keys}
            row_count = max((len(v) for v in lines_by_key.values()), default=1)
            manual_cell_indexes = {
                idx
                for idx, cell in enumerate(row.cells)
                if MANUAL_PLACEHOLDER_RE.search(" ".join(p.text for p in cell.paragraphs))
            }

            if row_count <= 1:
                # Нет размножения — обычная подстановка в текущую строку.
                _replace_manual_placeholders_in_row(table, i, manual, row_keys)
                i += 1
                continue

            # Вставляем row_count копий перед исходной строкой.
            for _ in range(row_count):
                clone = deepcopy(row._tr)
                row._tr.addprevious(clone)

            # Находим текущий индекс исходной (шаблонной) строки после вставки.
            original_idx = -1
            for idx2, r2 in enumerate(table.rows):
                if r2._tr is row._tr:
                    original_idx = idx2
                    break
            if original_idx == -1:
                i += 1
                continue

            # Заполняем вставленные строки построчно.
            start_idx = original_idx - row_count
            for offset in range(row_count):
                row_values = {
                    key: (vals[offset] if offset < len(vals) else "")
                    for key, vals in lines_by_key.items()
                }
                current_row_index = start_idx + offset
                _replace_manual_placeholders_in_row(table, current_row_index, row_values, row_keys)

                # Для добавленных строк (offset > 0) очищаем "статические" столбцы
                # без MANUAL-меток: так новые строки остаются в той же группе,
                # но не дублируют соседние колонки.
                if offset > 0 and manual_cell_indexes:
                    current_row = table.rows[current_row_index]
                    for cell_idx, cell in enumerate(current_row.cells):
                        if cell_idx in manual_cell_indexes:
                            continue
                        if _is_competency_cell_from_db(cell):
                            continue
                        _clear_cell_text_keep_paragraphs(cell)

            # Объединяем "статические" столбцы по вертикали на всю добавленную группу:
            # первая строка группы остаётся содержательной, следующие визуально
            # относятся к ней же и не создают самостоятельных записей.
            if row_count > 1 and manual_cell_indexes:
                top_row = table.rows[start_idx]
                for cell_idx in range(len(top_row.cells)):
                    if cell_idx in manual_cell_indexes:
                        continue
                    top_cell = table.rows[start_idx].cells[cell_idx]
                    for merge_row_idx in range(start_idx + 1, start_idx + row_count):
                        bottom_cell = table.rows[merge_row_idx].cells[cell_idx]
                        try:
                            top_cell = top_cell.merge(bottom_cell)
                        except Exception:
                            # Если merge невозможен из-за особенностей сетки таблицы,
                            # не останавливаем генерацию всего документа.
                            continue

            # Удаляем исходную шаблонную строку с плейсхолдерами.
            row._tr.getparent().remove(row._tr)

            # Продолжаем после вставленного диапазона.
            i = start_idx + row_count


def _force_run_times_new_roman(run) -> None:
    """Принудительно задаёт Times New Roman для run (включая кириллицу)."""
    run.font.name = "Times New Roman"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:cs"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "Times New Roman")


def enforce_times_new_roman(doc: Document) -> None:
    """Приводит весь текст документа к шрифту Times New Roman."""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            _force_run_times_new_roman(run)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        _force_run_times_new_roman(run)


def load_opop_data(json_path: str | Path = "opop_data.json") -> dict:
    """
    Загружает данные из JSON-файла.
    
    :param json_path: путь к JSON-файлу с данными
    :return: словарь с данными
    """
    json_path = Path(json_path)
    
    if not json_path.exists():
        raise FileNotFoundError(
            f"Не найден файл с данными: {json_path.resolve()}. "
            "Убедитесь, что файл opop_data.json существует."
        )
    
    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        return data
    except json.JSONDecodeError as e:
        raise ValueError(
            f"Ошибка при разборе JSON-файла: {json_path.resolve()}. "
            f"Проверьте корректность JSON-формата. Ошибка: {e}"
        ) from e


def parse_competencies_from_string(comp_string: str, prefix: str) -> list[dict]:
    """
    Преобразует строку с компетенциями в структурированный список.
    
    Формат строки: "УК-1 - Текст компетенции\nУК-2 - Текст компетенции"
    
    :param comp_string: строка с компетенциями
    :param prefix: префикс компетенций (УК, ОПК, ПК)
    :return: список словарей с компетенциями
    """
    competencies = []
    
    if not comp_string:
        return competencies
    
    # Разбиваем по строкам
    lines = comp_string.strip().split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Ищем разделитель " - " или "-"
        if ' - ' in line:
            code, description = line.split(' - ', 1)
        elif ' -' in line:
            code, description = line.split(' -', 1)
        elif '- ' in line:
            code, description = line.split('- ', 1)
        else:
            # Если разделитель не найден, используем всю строку как компетенцию
            code = line
            description = ""
        
        competencies.append({
            "category": "",  # Категорию можно заполнить позже при необходимости
            "competence": f"{code.strip()} - {description.strip()}" if description else code.strip(),
            "indicators": ""  # Индикаторы пока оставляем пустыми
        })
    
    return competencies


def _find_table_range(doc: Document, start_marker: str, end_marker: str):
    def _contains_marker(text: str, marker: str) -> bool:
        if marker in text:
            return True
        # Word иногда вносит лишние пробелы/переносы в тексте ячейки.
        norm_text = "".join(text.split())
        norm_marker = "".join(marker.split())
        return norm_marker in norm_text

    target_table = None
    template_row = None
    template_row_index = -1
    end_row_index = -1

    for table in doc.tables:
        for i, row in enumerate(table.rows):
            row_text = " ".join(cell.text for cell in row.cells)
            if _contains_marker(row_text, start_marker):
                target_table = table
                template_row = row
                template_row_index = i
                break
        if target_table:
            break

    if target_table is None or template_row is None:
        return None, None, -1, -1

    for i, row in enumerate(target_table.rows):
        row_text = " ".join(cell.text for cell in row.cells)
        if _contains_marker(row_text, end_marker):
            end_row_index = i
            break

    return target_table, template_row, template_row_index, end_row_index


def _replace_table_range_with_rows(
    target_table,
    template_row,
    template_row_index: int,
    end_row_index: int,
    rows_values: list[list[str]],
) -> None:
    # Удаляем маркеры из шаблонной строки
    for cell in template_row.cells:
        # Убираем любые START/END-маркеры, включая секционные (:06, :activity_1 и т.п.)
        cell.text = re.sub(r"\{\{(?:START_TABLE|END_TABLE):[^}]+\}\}", "", cell.text)

    if end_row_index == -1:
        end_row_index = template_row_index

    # Готовим новые строки
    new_rows = []
    for values in rows_values:
        new_row = deepcopy(template_row._tr)
        new_rows.append((new_row, values))

    # Удаляем шаблонный диапазон
    for i in range(end_row_index, template_row_index - 1, -1):
        if i < len(target_table.rows):
            row_to_remove = target_table.rows[i]
            row_to_remove._tr.getparent().remove(row_to_remove._tr)

    # Вставляем строки
    if not new_rows:
        return

    if template_row_index < len(target_table.rows):
        reference_row = target_table.rows[template_row_index]
        for row_element, _ in reversed(new_rows):
            reference_row._tr.addprevious(row_element)
    else:
        for row_element, _ in new_rows:
            target_table._tbl.append(row_element)

    # Заполняем значения
    start_idx = template_row_index
    for offset, (_, values) in enumerate(new_rows):
        row_idx = start_idx + offset
        if row_idx >= len(target_table.rows):
            break
        current_row = target_table.rows[row_idx]
        for col_idx, value in enumerate(values):
            if col_idx < len(current_row.cells):
                current_row.cells[col_idx].text = _sanitize_xml_text(value)


def _fill_table_section_rows(
    doc: Document,
    start_marker: str,
    end_marker: str,
    rows_values: list[list[str]],
) -> bool:
    target_table, template_row, template_row_index, end_row_index = _find_table_range(
        doc, start_marker, end_marker
    )
    if target_table is None or template_row is None:
        return False
    _replace_table_range_with_rows(
        target_table,
        template_row,
        template_row_index,
        end_row_index,
        rows_values,
    )
    return True


def _is_competencies_table_uk_opk(table) -> bool:
    """
    Эвристика: таблица УК/ОПК имеет 3 колонки и заголовки про категорию/компетенцию/индикатор.
    Нужна, чтобы не трогать другие таблицы при слиянии.
    """
    try:
        if not table.rows or len(table.rows[0].cells) < 3:
            return False
        header = " ".join(c.text for c in table.rows[0].cells[:3]).lower()
        return ("категор" in header) and ("компетенц" in header) and ("индикатор" in header)
    except Exception:
        return False


def _merge_vertical_category_groups(table, col_idx: int, start_row_idx: int) -> None:
    """
    Объединяет по вертикали ячейки столбца col_idx для таблиц УК/ОПК.

    Правила:
    - одинаковые подряд значения объединяются;
    - пустая ячейка объединяется с группой над ней (то есть "прилипает" вверх),
      даже если сверху непусто;
    - после merge в объединённой ячейке остаётся только верхний текст (без дублей).

    Доступ к ячейкам через table.cell(row, col): так корректна сетка при grid_before,
    объединениях и строках с разным числом «видимых» ячеек в row.cells.
    """
    if not table.rows or start_row_idx >= len(table.rows):
        return
    try:
        ncol = len(table.columns)
    except Exception:
        return
    if col_idx >= ncol:
        return

    def _cell_at(row_idx: int):
        try:
            return table.cell(row_idx, col_idx)
        except (IndexError, ValueError):
            return None

    group_start = start_row_idx
    top = _cell_at(start_row_idx)
    if top is None:
        return
    anchor_text = (top.text or "").strip()
    anchor_is_empty = not anchor_text

    for r in range(start_row_idx + 1, len(table.rows)):
        cell = _cell_at(r)
        if cell is None:
            continue
        text = (cell.text or "").strip()
        is_empty = not text

        # Пустые строки всегда относим к предыдущей группе.
        # Непустые — сливаем только если совпадают с anchor_text.
        should_merge = is_empty or (not anchor_is_empty and text == anchor_text) or (anchor_is_empty and is_empty)

        if should_merge:
            try:
                top_cell = table.cell(group_start, col_idx)
                top_text = (top_cell.text or "").strip()

                # Важно: python-docx при merge часто "склеивает" тексты.
                # Поэтому очищаем нижнюю ячейку перед merge и затем восстанавливаем верхний текст.
                _clear_cell_text_keep_paragraphs(cell)
                merged = top_cell.merge(cell)
                # В merged оставляем только верхний текст (без дублей)
                merged.text = top_text
            except Exception:
                pass
            continue

        # Начинаем новую группу
        group_start = r
        anchor_text = text
        anchor_is_empty = not anchor_text


def _merge_competencies_category_column_after_manual(doc: Document) -> None:
    """
    После подстановки manual_fields объединяет 1‑й столбец в таблицах УК/ОПК:
    - пустые ячейки сливаются с группой над ними,
    - одинаковые значения сливаются между собой,
    - в итоговой объединённой ячейке остаётся один текст (верхний).
    """
    for table in doc.tables:
        if not _is_competencies_table_uk_opk(table):
            continue
        # 0 — заголовок, 1.. — данные.
        _merge_vertical_category_groups(table, col_idx=0, start_row_idx=1)


def _is_pk_competencies_table(table) -> bool:
    """
    Эвристика: таблица ПК имеет минимум 5 столбцов и упоминает компетенции и индикаторы,
    но не является таблицей ОПК (общепрофессиональные).
    """
    try:
        if not table.rows:
            return False
        nc = len(table.columns)
        if nc < 5:
            return False
        header_parts = []
        for j in range(min(nc, 8)):
            try:
                header_parts.append(table.cell(0, j).text)
            except Exception:
                continue
        header = " ".join(header_parts).lower()
        if "общепрофессиональн" in header:
            return False
        return "индикатор" in header and "компетенц" in header
    except Exception:
        return False


def _merge_pk_manual_category_columns_after_manual(doc: Document) -> None:
    """
    После подстановки manual_fields: для таблицы ПК объединяет столбцы 1, 2 и 5
    (индексы 0, 1, 4) по тем же правилам, что категория в ОПК.
    Столбец 4 (индикаторы ПК) не трогаем — как третий столбец ОПК.
    """
    for table in doc.tables:
        try:
            if not _is_pk_competencies_table(table):
                continue
            n = len(table.columns)
            if n < 5:
                continue
            for col_idx in (0, 1, 4):
                if col_idx < n:
                    _merge_vertical_category_groups(table, col_idx=col_idx, start_row_idx=1)
        except Exception:
            continue


def replace_placeholders_in_paragraphs(doc: Document, data: dict[str, str]) -> None:
    """Замена меток в параграфах документа."""
    for paragraph in doc.paragraphs:
        if not paragraph.text:
            continue
        new_text = paragraph.text
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"  # из key делаем {{key}}
            if placeholder in new_text:
                new_text = new_text.replace(placeholder, _sanitize_xml_text(value))
        if new_text != paragraph.text:
            _set_paragraph_text_preserve_first_run_style(paragraph, new_text)


def replace_placeholders_in_tables(doc: Document, data: dict[str, str]) -> None:
    """Замена меток в ячейках таблиц документа."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if not paragraph.text:
                        continue
                    new_text = paragraph.text
                    for key, value in data.items():
                        placeholder = f"{{{{{key}}}}}"  # из key делаем {{key}}
                        if placeholder in new_text:
                            new_text = new_text.replace(placeholder, _sanitize_xml_text(value))
                    if new_text != paragraph.text:
                        _set_paragraph_text_preserve_first_run_style(paragraph, new_text)


def replace_placeholders(doc: Document, data: dict[str, str]) -> None:
    """Общая функция замены меток во всём документе."""
    replace_placeholders_in_paragraphs(doc, data)
    replace_placeholders_in_tables(doc, data)


def _decode_base64_image(data: str) -> bytes | None:
    """Декодирует подпись из raw base64 или data:image/...;base64,..."""
    if not data or not str(data).strip():
        return None
    s = str(data).strip()
    payload = s
    m = re.match(r"^data:image/[\w+.-]+;base64,(.+)$", s, re.IGNORECASE | re.DOTALL)
    if m:
        payload = m.group(1).strip()
    payload = re.sub(r"\s+", "", payload)
    if not payload:
        return None
    try:
        raw = base64.b64decode(payload, validate=False)
    except Exception:
        return None
    return raw if raw else None


def get_employee_signatures(opop_data: dict[str, Any]) -> dict[str, str]:
    """Подписи сотрудников: сначала employee_signatures, иначе _api_snapshot."""
    sigs = opop_data.get("employee_signatures")
    if isinstance(sigs, dict):
        return {str(k): str(v or "") for k, v in sigs.items()}
    snapshot = opop_data.get("_api_snapshot")
    if isinstance(snapshot, dict):
        nested = snapshot.get("employee_signatures")
        if isinstance(nested, dict):
            return {str(k): str(v or "") for k, v in nested.items()}
    return {}


def _clear_paragraph_runs(paragraph) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def _replace_image_placeholder_in_paragraph(
    paragraph,
    *,
    key: str,
    image_bytes: bytes,
    width_cm: float = _DEFAULT_SIGNATURE_WIDTH_CM,
) -> bool:
    placeholder = f"{{{{IMAGE:{key}}}}}"
    if placeholder not in (paragraph.text or ""):
        return False
    _clear_paragraph_runs(paragraph)
    paragraph.add_run().add_picture(BytesIO(image_bytes), width=Cm(width_cm))
    return True


def _replace_image_placeholders_in_paragraphs(
    paragraphs,
    signatures: dict[str, str],
    *,
    width_cm: float = _DEFAULT_SIGNATURE_WIDTH_CM,
) -> list[str]:
    """Возвращает ключи меток, для которых не удалось вставить изображение."""
    missing: list[str] = []
    for paragraph in paragraphs:
        text = paragraph.text or ""
        for m in IMAGE_PLACEHOLDER_RE.finditer(text):
            key = m.group(1)
            raw = signatures.get(key, "")
            image_bytes = _decode_base64_image(raw)
            if image_bytes:
                if not _replace_image_placeholder_in_paragraph(
                    paragraph, key=key, image_bytes=image_bytes, width_cm=width_cm
                ):
                    missing.append(key)
            else:
                # Нет данных — убираем метку, абзац оставляем пустым.
                placeholder = f"{{{{IMAGE:{key}}}}}"
                if placeholder in text:
                    _set_paragraph_text_preserve_first_run_style(
                        paragraph, text.replace(placeholder, "")
                    )
                missing.append(key)
    return missing


def replace_image_placeholders(
    doc: Document,
    opop_data: dict[str, Any],
    *,
    width_cm: float = _DEFAULT_SIGNATURE_WIDTH_CM,
) -> list[str]:
    """
    Заменяет {{IMAGE:key}} на картинку из employee_signatures[key] (base64).
    Возвращает ключи, для которых картинка не была вставлена.
    """
    signatures = get_employee_signatures(opop_data)
    if not signatures:
        return []

    missing: list[str] = []
    missing.extend(
        _replace_image_placeholders_in_paragraphs(
            doc.paragraphs, signatures, width_cm=width_cm
        )
    )
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                missing.extend(
                    _replace_image_placeholders_in_paragraphs(
                        cell.paragraphs, signatures, width_cm=width_cm
                    )
                )
    return sorted(set(missing))


def scan_image_template_keys(template_path: str | Path) -> list[str]:
    """Список ключей {{IMAGE:...}} в шаблоне (по XML, как для MANUAL)."""
    template_path = Path(template_path)
    if not template_path.exists():
        return []
    seen: set[str] = set()
    keys: list[str] = []
    try:
        with zipfile.ZipFile(template_path) as zf:
            for name in zf.namelist():
                if not name.startswith("word/") or not name.endswith(".xml"):
                    continue
                if "media" in name:
                    continue
                try:
                    xml = zf.read(name).decode("utf-8")
                except Exception:
                    continue
                for m in IMAGE_PLACEHOLDER_RE.finditer(xml):
                    k = m.group(1)
                    if k not in seen:
                        seen.add(k)
                        keys.append(k)
    except (OSError, zipfile.BadZipFile):
        return []
    return keys


def scan_manual_template_keys(template_path: str | Path) -> list[str]:
    """
    Находит в шаблоне все метки вида {{MANUAL:key}} по XML Word (надёжнее, чем только python-docx).
    Порядок — по первому появлению в XML.
    """
    template_path = Path(template_path)
    if not template_path.exists():
        return []
    seen: set[str] = set()
    keys: list[str] = []
    try:
        with zipfile.ZipFile(template_path) as zf:
            for name in zf.namelist():
                if not name.startswith("word/") or not name.endswith(".xml"):
                    continue
                if "media" in name:
                    continue
                try:
                    xml = zf.read(name).decode("utf-8")
                except Exception:
                    continue
                for m in MANUAL_PLACEHOLDER_RE.finditer(xml):
                    k = m.group(1)
                    if k not in seen:
                        seen.add(k)
                        keys.append(k)
    except (OSError, zipfile.BadZipFile):
        return []
    return keys


def replace_manual_placeholders_in_paragraphs(doc: Document, manual: dict[str, str]) -> None:
    for paragraph in doc.paragraphs:
        if not paragraph.text:
            continue
            
        original_text = paragraph.text
        new_text = original_text
        
        # Проверяем наличие меток
        has_placeholder = False
        for key, value in manual.items():
            placeholder = f"{{{{MANUAL:{key}}}}}"
            if placeholder in new_text:
                new_text = new_text.replace(placeholder, _normalize_manual_text(value or ""))
                has_placeholder = True
        
        if not has_placeholder:
            continue
            
        # Если текст изменился, обновляем текущий абзац с сохранением формата.
        if new_text != original_text:
            paragraph.paragraph_format.first_line_indent = Pt(0)
            _set_paragraph_multiline_text_preserve_style(paragraph, new_text)


def replace_manual_placeholders_in_tables(doc: Document, manual: dict[str, str]) -> None:
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if not paragraph.text:
                        continue
                    
                    original_text = paragraph.text
                    new_text = original_text
                    
                    has_placeholder = False
                    for key, value in manual.items():
                        placeholder = f"{{{{MANUAL:{key}}}}}"
                        if placeholder in new_text:
                            new_text = new_text.replace(placeholder, _normalize_manual_text(value or ""))
                            has_placeholder = True
                    
                    if not has_placeholder:
                        continue
                    
                    if new_text != original_text:
                        paragraph.paragraph_format.first_line_indent = Pt(0)
                        _set_paragraph_multiline_text_preserve_style(paragraph, new_text)


def replace_manual_placeholders(doc: Document, manual: dict[str, Any]) -> None:
    if not isinstance(manual, dict):
        return
    flat: dict[str, str] = {
        str(k): _sanitize_xml_text(v) for k, v in manual.items()
    }
    _expand_manual_rows_in_tables(doc, flat)
    replace_manual_placeholders_in_paragraphs(doc, flat)
    replace_manual_placeholders_in_tables(doc, flat)


def fill_competencies_table(
    doc: Document,
    start_marker: str,
    end_marker: str,
    competencies_data: list[dict],
    *,
    manual_prefix: str | None = None,
) -> None:
    """
    Универсальная функция заполнения таблицы компетенций.
    
    :param doc: документ Word
    :param start_marker: маркер начала таблицы (например, {{START_TABLE:universal}})
    :param end_marker: маркер конца таблицы (например, {{END_TABLE:universal}})
    :param competencies_data: список словарей с данными компетенций
    """
    target_table, template_row, template_row_index, end_row_index = _find_table_range(
        doc, start_marker, end_marker
    )
    if target_table is None or template_row is None:
        print(f"Предупреждение: не найдена таблица с маркером {start_marker}")
        return

    rows_values = []
    for idx, comp_data in enumerate(competencies_data, start=1):
        indicators_value = str(comp_data.get("indicators", ""))
        if manual_prefix:
            indicators_value = f"{{{{MANUAL:{manual_prefix}_{idx}_ind}}}}"

        # 1‑й столбец: ручное поле категории, если задан префикс.
        # Это позволяет при необходимости ввести одинаковый текст категории
        # для нескольких компетенций (просто повторив значение в разных полях).
        if manual_prefix:
            category_value = f"{{{{MANUAL:{manual_prefix}_{idx}_cat}}}}"
        else:
            category_value = str(comp_data.get("category", ""))

        rows_values.append(
            [
                category_value,
                str(comp_data.get("competence", "")),
                indicators_value,
            ]
        )

    _replace_table_range_with_rows(
        target_table,
        template_row,
        template_row_index,
        end_row_index,
        rows_values,
    )


def fill_universal_competencies_table(doc: Document, competencies_string: str) -> None:
    """Заполняет таблицу универсальных компетенций."""
    competencies_data = parse_competencies_from_string(competencies_string, "УК")
    fill_competencies_table(
        doc,
        start_marker="{{START_TABLE:universal}}",
        end_marker="{{END_TABLE:universal}}",
        competencies_data=competencies_data,
        manual_prefix="uk",
    )


def fill_professional_competencies_table(doc: Document, competencies_string: str) -> None:
    """Заполняет таблицу общепрофессиональных компетенций (ОПК)."""
    competencies_data = parse_competencies_from_string(competencies_string, "ОПК")
    fill_competencies_table(
        doc,
        start_marker="{{START_TABLE:professional}}",
        end_marker="{{END_TABLE:professional}}",
        competencies_data=competencies_data,
        manual_prefix="opk",
    )


def _parse_prof_standards_table(opop_data: dict[str, Any]) -> list[dict[str, str]]:
    # Новый формат (предпочтительный): массив объектов
    raw = opop_data.get("prof_standards_table")
    if isinstance(raw, list):
        rows = []
        for item in raw:
            if not isinstance(item, dict):
                continue
            area = str(item.get("area", ""))
            area_code = str(item.get("area_code", "")).strip()
            if not area_code:
                if area.strip().startswith("06"):
                    area_code = "06"
                elif area.strip().startswith("40"):
                    area_code = "40"
            rows.append(
                {
                    "area_code": area_code,
                    "area": area,
                    "standard": str(item.get("standard", "")),
                    "generalized_functions": str(item.get("generalized_functions", "")),
                }
            )
        if rows:
            return rows

    # Старый формат: area_06/area_40 + prof_standard_1..N (fallback-заглушка)
    areas = [str(opop_data.get("area_06", "")), str(opop_data.get("area_40", ""))]
    standards = [
        str(opop_data.get("prof_standard_1", "")),
        str(opop_data.get("prof_standard_2", "")),
        str(opop_data.get("prof_standard_3", "")),
    ]
    rows = []
    if areas[0]:
        for s in standards[:2]:
            if s:
                rows.append({"area": areas[0], "standard": s})
    if areas[1] and standards[2]:
        rows.append({"area": areas[1], "standard": standards[2]})
    for r in rows:
        if "06 " in r.get("area", ""):
            r["area_code"] = "06"
        elif "40 " in r.get("area", ""):
            r["area_code"] = "40"
        else:
            r["area_code"] = ""
        r.setdefault("generalized_functions", "")
    return rows


def fill_prof_standards_table(doc: Document, opop_data: dict[str, Any]) -> None:
    rows = _parse_prof_standards_table(opop_data)
    by_area: dict[str, list[dict[str, str]]] = {}
    for r in rows:
        by_area.setdefault(r.get("area_code", ""), []).append(r)

    # Новый режим: отдельные секции по коду области (как на вашем шаблоне 06/40).
    # В секциях заполняются 1-2 столбцы: профстандарт + обобщенные трудовые функции.
    used_sectional = False
    for area_code, items in by_area.items():
        if not area_code:
            continue
        start = f"{{{{START_TABLE:prof_standards:{area_code}}}}}"
        end = f"{{{{END_TABLE:prof_standards:{area_code}}}}}"
        section_rows: list[list[str]] = []
        # Для каждого профстандарта создаём свои manual-ключи, чтобы можно было
        # вводить "трудовые функции" и "уровень квалификации" раздельно.
        for idx, item in enumerate(items, start=1):
            tf_key = f"ps_{area_code}_{idx}_tf"
            lvl_key = f"ps_{area_code}_{idx}_lvl"
            section_rows.append(
                [
                    item.get("standard", ""),
                    item.get("generalized_functions", ""),
                    f"{{{{MANUAL:{tf_key}}}}}",
                    f"{{{{MANUAL:{lvl_key}}}}}",
                ]
            )
        ok = _fill_table_section_rows(doc, start, end, section_rows)
        used_sectional = used_sectional or ok

    if used_sectional:
        return

    # Совместимость со старым единым маркером.
    start_marker = "{{START_TABLE:prof_standards}}"
    end_marker = "{{END_TABLE:prof_standards}}"
    rows_values = [[r.get("area", ""), r.get("standard", "")] for r in rows]
    if not _fill_table_section_rows(doc, start_marker, end_marker, rows_values):
        print(f"Предупреждение: не найдена таблица с маркерами {start_marker} / {end_marker}")


def _parse_pk_table(opop_data: dict[str, Any]) -> list[dict[str, str]]:
    # Новый формат (предпочтительный):
    # [
    #   {"task_type":"...", "pk_code":"ПК-1", "pk_description":"...", "indicators":"..."},
    #   ...
    # ]
    raw = opop_data.get("pk_table")
    if isinstance(raw, list):
        rows = []
        for item in raw:
            if not isinstance(item, dict):
                continue
            rows.append(
                {
                    "task_type_key": str(item.get("task_type_key", "")),
                    "task_type": str(item.get("task_type", "")),
                    "pk_code": str(item.get("pk_code", "")),
                    "pk_description": str(item.get("pk_description", "")),
                    "indicators": str(item.get("indicators", "")),
                }
            )
        if rows:
            return rows

    # Fallback: из текстовой строки professional_competencies
    # В этом случае группировка по типам задач недоступна, всё уходит в "Не указан тип".
    fallback = parse_competencies_from_string(str(opop_data.get("professional_competencies", "")), "ПК")
    rows = []
    for c in fallback:
        comp = str(c.get("competence", ""))
        code = comp.split(" - ", 1)[0].strip() if comp else ""
        desc = comp.split(" - ", 1)[1].strip() if " - " in comp else ""
        rows.append(
            {
                "task_type_key": "",
                "task_type": "Не указан тип задач профессиональной деятельности",
                "pk_code": code,
                "pk_description": desc,
                "indicators": "",
            }
        )
    return rows


def fill_pk_table(doc: Document, opop_data: dict[str, Any]) -> None:
    rows = _parse_pk_table(opop_data)
    # Если task_type_key не задан, пробуем вывести из activity_1..activity_3.
    activity_map = {
        str(opop_data.get("activity_1", "")).strip(): "activity_1",
        str(opop_data.get("activity_2", "")).strip(): "activity_2",
        str(opop_data.get("activity_3", "")).strip(): "activity_3",
    }
    for r in rows:
        if not r.get("task_type_key"):
            r["task_type_key"] = activity_map.get(str(r.get("task_type", "")).strip(), "")

    # Новый режим: отдельные секции PK под каждым типом задач.
    # col1,col2,col5 — MANUAL как категории ОПК; col3 из БД; col4 — MANUAL как индикаторы ОПК.
    grouped: dict[str, list[dict[str, str]]] = {}
    for r in rows:
        key = r.get("task_type_key", "")
        grouped.setdefault(key, []).append(r)

    used_sectional = False
    for key, items in grouped.items():
        if not key:
            continue
        start = f"{{{{START_TABLE:pk:{key}}}}}"
        end = f"{{{{END_TABLE:pk:{key}}}}}"
        section_rows: list[list[str]] = []
        for i in items:
            code = i.get("pk_code", "")
            desc = i.get("pk_description", "")
            pk_text = f"{code} - {desc}".strip(" -")
            row_idx = len(section_rows) + 1
            section_rows.append(
                [
                    f"{{{{MANUAL:pk_{key}_{row_idx}_cat1}}}}",
                    f"{{{{MANUAL:pk_{key}_{row_idx}_cat2}}}}",
                    pk_text,
                    f"{{{{MANUAL:pk_{key}_{row_idx}_ind}}}}",
                    f"{{{{MANUAL:pk_{key}_{row_idx}_cat5}}}}",
                ]
            )
        ok = _fill_table_section_rows(doc, start, end, section_rows)
        used_sectional = used_sectional or ok

    if used_sectional:
        return

    # Совместимость со старым единым маркером (шаблон с 5 столбцами).
    start_marker = "{{START_TABLE:pk}}"
    end_marker = "{{END_TABLE:pk}}"
    rows_values: list[list[str]] = []
    for r in rows:
        code = r.get("pk_code", "")
        desc = r.get("pk_description", "")
        pk_text = f"{code} - {desc}".strip(" -")
        row_idx = len(rows_values) + 1
        rows_values.append(
            [
                f"{{{{MANUAL:pk_{row_idx}_cat1}}}}",
                f"{{{{MANUAL:pk_{row_idx}_cat2}}}}",
                pk_text,
                f"{{{{MANUAL:pk_{row_idx}_ind}}}}",
                f"{{{{MANUAL:pk_{row_idx}_cat5}}}}",
            ]
        )

    if not _fill_table_section_rows(doc, start_marker, end_marker, rows_values):
        print(f"Предупреждение: не найдена таблица с маркерами {start_marker} / {end_marker}")


def generate_opop_document(
    template_path: str | Path = "template.docx",
    data_path: str | Path = "opop_data.json",
    output_path: str | Path = "ОПОП-ПМ_тестовый.docx",
    *,
    skip_manual_replace: bool = False,
) -> Path:
    """
    Создаёт документ ОПОП на основе шаблона и данных из JSON.
    
    :param template_path: путь к шаблону Word (.docx)
    :param data_path: путь к JSON-файлу с данными
    :param output_path: путь для сохранения сгенерированного документа
    :return: путь к созданному файлу
    """
    template_path = Path(template_path)
    data_path = Path(data_path)
    output_path = Path(output_path)
    
    # Проверка существования файлов
    if not template_path.exists():
        raise FileNotFoundError(
            f"Не найден файл шаблона: {template_path.resolve()}. "
            "Создайте template.docx и добавьте в него метки вида {{direction_code}} и т.п."
        )
    
    # Загружаем данные из JSON
    opop_data = load_opop_data(data_path)
    
    # Проверка, что файл шаблона является корректным .docx
    with template_path.open("rb") as f:
        signature = f.read(8)
    if not signature.startswith(b"PK"):
        raise ValueError(
            "Файл шаблона не является корректным .docx (zip-архивом). "
            f"Путь: {template_path.resolve()}. "
            "Откройте шаблон в Word и выполните 'Сохранить как' → 'Документ Word (*.docx)'. "
            "Если у вас сейчас формат .doc, переименования расширения недостаточно."
        )
    
    try:
        doc = Document(template_path)
    except BadZipFile as e:
        raise ValueError(
            "Не удалось открыть шаблон как .docx (файл повреждён или не является zip-архивом). "
            f"Путь: {template_path.resolve()}."
        ) from e
    
    # 1. Заменяем простые метки (все строковые значения из JSON)
    # Исключаем поля, которые используются для динамических таблиц
    excluded_keys = {
        "universal_competencies",
        "opk_competencies",
        "professional_competencies",
        "prof_standards_table",
        "pk_table",
        "manual_fields",
        "employee_signatures",
        "_api_snapshot",
        "_manual_fields_seed",
        "_missing_fields",
        "_missing_fields_text",
        "_source_export_time",
        "_generated_at",
    }
    simple_data = {
        k: _sanitize_xml_text(v)
        for k, v in opop_data.items()
        if k not in excluded_keys and isinstance(v, str)
    }
    replace_placeholders(doc, simple_data)

    missing_images = replace_image_placeholders(doc, opop_data)
    if missing_images:
        print(
            "Предупреждение: не вставлены подписи для меток IMAGE: "
            + ", ".join(missing_images)
        )
    
    # 2. Заполняем динамические таблицы
    if "universal_competencies" in opop_data:
        fill_universal_competencies_table(doc, opop_data["universal_competencies"])
    else:
        print("Предупреждение: в JSON-файле не найдены универсальные компетенции")
    
    if "opk_competencies" in opop_data:
        fill_professional_competencies_table(doc, opop_data["opk_competencies"])
    else:
        print("Предупреждение: в JSON-файле не найдены общепрофессиональные компетенции")

    fill_prof_standards_table(doc, opop_data)
    fill_pk_table(doc, opop_data)

    # Ручные поля: {{MANUAL:key}} — подставляем после таблиц (можно отключить для черновика).
    if not skip_manual_replace:
        replace_manual_placeholders(doc, opop_data.get("manual_fields", {}))
        _merge_competencies_category_column_after_manual(doc)
        _merge_pk_manual_category_columns_after_manual(doc)

    # Единый шрифт для всего документа: Times New Roman.
    enforce_times_new_roman(doc)

    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    try:
        result = generate_opop_document()
        print(f"Документ успешно сгенерирован: {result.resolve()}")
    except Exception as e:
        print(f"Ошибка при генерации документа: {e}")